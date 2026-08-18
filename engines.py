"""Conversion engines for Converter.

Each family (image / video / audio / document / pdf / archive) maps to a set of
valid target formats and a convert function. External tools (ffmpeg, pandoc,
Edge headless) are located lazily so the GUI can open even if one is missing -
the error surfaces on the actual conversion with an instructive message.
"""

import gzip
import os
import shutil
import subprocess
import sys
import tempfile
import zipfile
import tarfile
from pathlib import Path

CREATE_NO_WINDOW = 0x08000000  # keep console-less: we run under pythonw.exe

IMAGE_EXTS = {"png", "jpg", "jpeg", "webp", "bmp", "tif", "tiff", "gif", "ico",
              "heic", "heif", "avif"}
VIDEO_EXTS = {"mp4", "mkv", "mov", "avi", "webm", "wmv", "flv", "m4v", "mpg",
              "mpeg", "ts", "3gp"}
AUDIO_EXTS = {"mp3", "wav", "flac", "aac", "ogg", "oga", "m4a", "wma", "opus",
              "aiff", "aif"}
DOC_EXTS = {"docx", "md", "markdown", "html", "htm", "odt", "rtf", "epub",
            "txt", "tex"}
# Compound extensions must be checked before the bare single-suffix ones.
COMPOUND_EXTS = ("tar.gz", "tar.bz2", "tar.xz")
ARCHIVE_EXTS = {"zip", "7z", "tar", "gz", "tgz", "bz2", "xz", "tar.gz",
                "tar.bz2", "tar.xz"}
# Single-file (non-container) compression formats: repack decompresses first.
SINGLE_FILE_COMPRESSION = {"gz", "bz2", "xz"}

LOSSY_IMAGE_TARGETS = {"jpg", "webp", "heic", "avif"}
# Targets that can hold multiple frames, so animation is preserved.
ANIMATED_TARGETS = {"webp", "gif"}


class ConvertError(Exception):
    pass


# ---------------------------------------------------------------- detection

def ext_of(path):
    name = Path(path).name.lower()
    for multi in COMPOUND_EXTS:
        if name.endswith("." + multi):
            return multi
    return Path(name).suffix.lstrip(".")


def _split_name(path):
    """Split a filename into (base, dotext), honoring compound extensions.

    dotext keeps its leading dot (or is "" when there is no extension), so
    base + dotext == the original name. "backup.tar.gz" -> ("backup",
    ".tar.gz"); "Photo.PNG" -> ("Photo", ".PNG").
    """
    name = Path(path).name
    low = name.lower()
    for multi in COMPOUND_EXTS:
        if low.endswith("." + multi):
            cut = len(multi) + 1
            return name[:-cut], name[-cut:]
    suffix = Path(name).suffix
    return (name[:-len(suffix)] if suffix else name), suffix


def family_of(path):
    if os.path.isdir(path):
        return "folder"
    ext = ext_of(path)
    if ext in IMAGE_EXTS:
        return "image"
    if ext in VIDEO_EXTS:
        return "video"
    if ext in AUDIO_EXTS:
        return "audio"
    if ext == "pdf":
        return "pdf"
    if ext in DOC_EXTS:
        return "document"
    if ext in ARCHIVE_EXTS:
        return "archive"
    return "other"


# Targets per family: list of (label, target_ext). "conv" targets only.
FAMILY_TARGETS = {
    "image": [("PNG", "png"), ("JPG", "jpg"), ("WEBP", "webp"), ("GIF", "gif"),
              ("BMP", "bmp"), ("TIFF", "tiff"), ("ICO", "ico"),
              ("HEIC", "heic"), ("PDF", "pdf")],
    "video": [("MP4", "mp4"), ("MKV", "mkv"), ("MOV", "mov"), ("AVI", "avi"),
              ("WEBM", "webm"), ("GIF", "gif"), ("MP3 (audio)", "mp3"),
              ("WAV (audio)", "wav"), ("FLAC (audio)", "flac"),
              ("M4A (audio)", "m4a")],
    "audio": [("MP3", "mp3"), ("WAV", "wav"), ("FLAC", "flac"),
              ("M4A", "m4a"), ("OGG", "ogg"), ("OPUS", "opus"),
              ("AAC", "aac")],
    "document": [("DOCX", "docx"), ("PDF", "pdf"), ("Markdown", "md"),
                 ("HTML", "html"), ("ODT", "odt"), ("RTF", "rtf"),
                 ("EPUB", "epub"), ("TXT", "txt")],
    "pdf": [("DOCX", "docx"), ("TXT", "txt"), ("PNG (pages)", "png"),
            ("JPG (pages)", "jpg")],
}

COMPRESS_TARGETS = [("ZIP", "zip"), ("7Z", "7z"), ("TAR.GZ", "tar.gz")]


def targets_for(paths):
    """Return the choice list for a selection: (kind, value, label) tuples.

    kind: 'conv' converts each file, 'compress' packs the whole selection
    into one archive, 'extract' unpacks archives.
    """
    families = {family_of(p) for p in paths}
    choices = []

    if families == {"archive"}:
        choices.append(("extract", "", "Extract here"))
        own = {ext_of(p) for p in paths} if len(paths) == 1 else set()
        for label, ext in COMPRESS_TARGETS:
            if ext not in own:
                choices.append(("conv", ext, label + " (repack)"))
        return choices

    if len(families) == 1:
        fam = next(iter(families))
        targets = FAMILY_TARGETS.get(fam, [])
        own = {ext_of(p) for p in paths} if len(paths) == 1 else set()
        if "jpeg" in own:
            own.add("jpg")
        for label, ext in targets:
            if ext not in own:
                choices.append(("conv", ext, label))
    else:
        # Mixed selection: offer only targets valid for every file.
        common = None
        for fam in families:
            exts = {e for _, e in FAMILY_TARGETS.get(fam, [])}
            common = exts if common is None else common & exts
        for fam in families:
            for label, ext in FAMILY_TARGETS.get(fam, []):
                if ext in (common or set()) and not any(c[1] == ext for c in choices):
                    choices.append(("conv", ext, label))

    for label, ext in COMPRESS_TARGETS:
        choices.append(("compress", ext, label + " (compress)"))
    return choices


# ---------------------------------------------------------------- helpers

def unique_path(path):
    """A non-colliding path, inserting ' (n)' before the (compound) extension."""
    path = Path(path)
    if not path.exists():
        return path
    base, ext = _split_name(path)
    for i in range(1, 10000):
        cand = path.with_name(f"{base} ({i}){ext}")
        if not cand.exists():
            return cand
    raise ConvertError(f"Could not find a free name for {path}")


def _run(cmd, tool, cwd=None):
    proc = subprocess.run(cmd, capture_output=True, text=True,
                          encoding="utf-8", errors="replace", cwd=cwd,
                          creationflags=CREATE_NO_WINDOW)
    if proc.returncode != 0:
        tail = (proc.stderr or proc.stdout or "").strip().splitlines()[-8:]
        raise ConvertError(f"{tool} failed:\n" + "\n".join(tail))
    return proc


def _run_producing(cmd, tool, dst, cwd=None):
    """Run a tool that writes `dst`; delete a partial `dst` if it fails."""
    try:
        return _run(cmd, tool, cwd=cwd)
    except Exception:
        Path(dst).unlink(missing_ok=True)
        raise


def _app_dir():
    """Folder the app runs from (the exe's dir when frozen, else this file's)."""
    if getattr(sys, "frozen", False):
        return os.path.dirname(sys.executable)
    return os.path.dirname(os.path.abspath(__file__))


def _bundled(name):
    """A tool shipped alongside the app (same dir or a bin/ subfolder)."""
    d = _app_dir()
    for cand in (os.path.join(d, name), os.path.join(d, "bin", name)):
        if os.path.isfile(cand):
            return cand
    return None


def _find_tool(name, extra_candidates=()):
    found = shutil.which(name)
    if found:
        return found
    for cand in extra_candidates:
        if os.path.isfile(cand):
            return cand
    return None


def _import_py7zr():
    try:
        import py7zr
        return py7zr
    except ImportError:
        raise ConvertError("py7zr not installed. "
                           "Run: pip install --user py7zr")


def _ffmpeg():
    ff = _bundled("ffmpeg.exe") or _find_tool("ffmpeg")
    if not ff:
        raise ConvertError("FFmpeg not found. Install it with:\n"
                           "winget install Gyan.FFmpeg")
    return ff


def _pandoc():
    local = os.environ.get("LOCALAPPDATA", "")
    pd = _bundled("pandoc.exe") or _find_tool("pandoc", (
        os.path.join(local, "Pandoc", "pandoc.exe"),
        r"C:\Program Files\Pandoc\pandoc.exe",
    ))
    if not pd:
        raise ConvertError("Pandoc not found (needed for document conversion)."
                           "\nInstall it with: winget install JohnMacFarlane.Pandoc")
    return pd


def _edge():
    edge = _find_tool("msedge", (
        r"C:\Program Files (x86)\Microsoft\Edge\Application\msedge.exe",
        r"C:\Program Files\Microsoft\Edge\Application\msedge.exe",
    ))
    if not edge:
        raise ConvertError("Microsoft Edge not found (used for PDF export).")
    return edge


# ---------------------------------------------------------------- images

def _open_image(src):
    try:
        from PIL import Image
    except ImportError:
        raise ConvertError("Pillow not installed. Run: pip install --user pillow")
    if ext_of(src) in ("heic", "heif", "avif"):
        try:
            import pillow_heif
            pillow_heif.register_heif_opener()
            # AVIF moved out of pillow-heif in 1.x; Pillow 11.3+ reads it
            # natively, so only register the legacy opener if it exists.
            if hasattr(pillow_heif, "register_avif_opener"):
                pillow_heif.register_avif_opener()
        except ImportError:
            raise ConvertError("pillow-heif not installed. "
                               "Run: pip install --user pillow-heif")
    try:
        img = Image.open(src)
        img.load()
    except ConvertError:
        raise
    except Exception as e:
        raise ConvertError(f"Could not read image {Path(src).name}: "
                           f"{e.__class__.__name__}")
    return img


def _register_heif_save():
    try:
        import pillow_heif
        pillow_heif.register_heif_opener()
    except ImportError:
        raise ConvertError("pillow-heif not installed. "
                           "Run: pip install --user pillow-heif")


def _flatten(img):
    """Composite transparency onto white for formats without alpha."""
    from PIL import Image
    if img.mode in ("RGBA", "LA", "PA") or (img.mode == "P" and
                                            "transparency" in img.info):
        img = img.convert("RGBA")
        bg = Image.new("RGB", img.size, (255, 255, 255))
        bg.paste(img, mask=img.split()[-1])
        return bg
    return img.convert("RGB") if img.mode not in ("RGB", "L") else img


def convert_image(src, dst_ext, quality=85):
    img = _open_image(src)
    dst = unique_path(Path(src).with_suffix("." + dst_ext))
    frames = getattr(img, "n_frames", 1)
    animated = frames > 1 and dst_ext in ANIMATED_TARGETS

    if dst_ext in ("jpg", "bmp", "pdf"):
        img = _flatten(img)

    if dst_ext == "jpg":
        img.save(dst, "JPEG", quality=quality, optimize=True)
    elif dst_ext == "webp":
        img.save(dst, "WEBP", quality=quality, save_all=animated)
    elif dst_ext == "gif":
        img.save(dst, "GIF", save_all=animated)
    elif dst_ext in ("heic", "avif"):
        _register_heif_save()
        img.save(dst, quality=quality)
    elif dst_ext == "ico":
        img.save(dst, "ICO", sizes=[(16, 16), (32, 32), (48, 48), (256, 256)])
    elif dst_ext == "pdf":
        img.save(dst, "PDF", resolution=96)
    elif dst_ext in ("tif", "tiff"):
        img.save(dst, "TIFF")
    else:
        img.save(dst)
    return [str(dst)]


def images_to_pdf(paths, quality=85):
    paths = [os.path.abspath(p) for p in paths]
    pages = [_flatten(_open_image(p)) for p in paths]
    dst = unique_path(Path(paths[0]).with_suffix(".pdf"))
    pages[0].save(dst, "PDF", resolution=96, save_all=True,
                  append_images=pages[1:])
    return [str(dst)]


# ---------------------------------------------------------------- video/audio

# Force even dimensions: libx264 with 4:2:0 rejects odd width/height.
_EVEN = "scale=trunc(iw/2)*2:trunc(ih/2)*2"

_VIDEO_ENCODE = {
    "mp4": ["-c:v", "libx264", "-crf", "20", "-preset", "veryfast",
            "-vf", _EVEN, "-pix_fmt", "yuv420p", "-c:a", "aac", "-b:a", "192k"],
    "mkv": ["-c:v", "libx264", "-crf", "20", "-preset", "veryfast",
            "-vf", _EVEN, "-c:a", "aac", "-b:a", "192k"],
    "mov": ["-c:v", "libx264", "-crf", "20", "-preset", "veryfast",
            "-vf", _EVEN, "-pix_fmt", "yuv420p", "-c:a", "aac", "-b:a", "192k"],
    "avi": ["-c:v", "libx264", "-crf", "20", "-preset", "veryfast",
            "-vf", _EVEN, "-c:a", "libmp3lame", "-q:a", "3"],
    "webm": ["-c:v", "libvpx-vp9", "-crf", "33", "-b:v", "0", "-row-mt", "1",
             "-c:a", "libopus", "-b:a", "128k"],
}

_AUDIO_ENCODE = {
    "mp3": ["-c:a", "libmp3lame", "-q:a", "2"],
    "wav": ["-c:a", "pcm_s16le"],
    "flac": ["-c:a", "flac"],
    "m4a": ["-c:a", "aac", "-b:a", "192k"],
    "aac": ["-c:a", "aac", "-b:a", "192k"],
    "ogg": ["-c:a", "libvorbis", "-q:a", "5"],
    "opus": ["-c:a", "libopus", "-b:a", "128k"],
}

_REMUX_OK = {"mp4", "mkv", "mov", "m4v"}  # try stream-copy first for these


def convert_av(src, dst_ext):
    ff = _ffmpeg()
    dst = unique_path(Path(src).with_suffix("." + dst_ext))
    base = [ff, "-y", "-hide_banner", "-loglevel", "error", "-i", src]

    if dst_ext in _AUDIO_ENCODE:
        _run_producing(base + ["-vn"] + _AUDIO_ENCODE[dst_ext] + [str(dst)],
                       "FFmpeg", dst)
        return [str(dst)]

    if dst_ext == "gif":
        filt = ("fps=12,scale='min(480,iw)':-2:flags=lanczos,"
                "split[a][b];[a]palettegen[p];[b][p]paletteuse")
        _run_producing(base + ["-filter_complex", filt, str(dst)],
                       "FFmpeg", dst)
        return [str(dst)]

    if dst_ext in _REMUX_OK and ext_of(src) in _REMUX_OK:
        # Same codec families - a container swap is instant and lossless.
        proc = subprocess.run(base + ["-c", "copy", str(dst)],
                              capture_output=True, text=True,
                              encoding="utf-8", errors="replace",
                              creationflags=CREATE_NO_WINDOW)
        if proc.returncode == 0:
            return [str(dst)]
        Path(dst).unlink(missing_ok=True)

    _run_producing(base + _VIDEO_ENCODE[dst_ext] + [str(dst)], "FFmpeg", dst)
    return [str(dst)]


# ---------------------------------------------------------------- documents

_PANDOC_READERS = {"txt": "markdown", "md": "markdown", "markdown": "markdown",
                   "htm": "html", "tex": "latex"}
# Writers that need explicit -s; others (docx/odt/epub/pptx) imply standalone.
_STANDALONE_TARGETS = {"html", "epub", "rtf"}


def _pandoc_args(src):
    ext = ext_of(src)
    reader = _PANDOC_READERS.get(ext)
    args = ["-f", reader] if reader else []
    # Resolve relative resources (images) against the source's own folder,
    # not the process working directory.
    args += ["--resource-path", str(Path(src).parent)]
    return args


def convert_doc(src, dst_ext):
    pd = _pandoc()
    if dst_ext == "pdf":
        return _doc_to_pdf(src)
    src = os.path.abspath(src)
    dst = unique_path(Path(src).with_suffix("." + dst_ext))
    args = [pd, str(src)] + _pandoc_args(src)
    if dst_ext == "txt":
        args += ["-t", "plain"]
    if dst_ext in _STANDALONE_TARGETS:
        args += ["-s"]
    _run_producing(args + ["-o", str(dst)], "Pandoc", dst)
    return [str(dst)]


# Office formats carry real fonts, styles, alignment and page layout, so they
# must be rendered by an actual office engine. Flattening them through pandoc
# (as markdown/html are) discards all of that formatting.
_OFFICE_TO_PDF = {"docx", "odt", "rtf"}


def _doc_to_pdf(src):
    """docx/odt/rtf -> PDF via a real office renderer (preserves formatting);
    markdown/html/txt/tex/epub -> pandoc + Edge (no rich layout to keep)."""
    src = os.path.abspath(src)
    if ext_of(src) in _OFFICE_TO_PDF:
        for engine in (_office_pdf_word, _office_pdf_soffice):
            try:
                return engine(src)
            except ConvertError:
                continue  # engine unavailable - try the next one
        # Neither Word nor LibreOffice available: fall back to pandoc so the
        # user still gets a PDF (lower fidelity - documented in the README).
    return _doc_to_pdf_pandoc(src)


def _office_pdf_word(src):
    """Export through Microsoft Word (best fidelity) if Word is installed."""
    try:
        import pythoncom
        import win32com.client as win32
    except ImportError:
        raise ConvertError("Word automation unavailable (pywin32 missing).")
    dst = unique_path(Path(src).with_suffix(".pdf"))
    pythoncom.CoInitialize()
    word = doc = None
    try:
        try:
            word = win32.DispatchEx("Word.Application")
        except Exception:
            raise ConvertError("Microsoft Word is not installed.")
        word.Visible = False
        try:
            word.DisplayAlerts = 0
        except Exception:
            pass
        doc = word.Documents.Open(str(src), ReadOnly=True,
                                  AddToRecentFiles=False)
        doc.ExportAsFixedFormat(str(dst), 17)  # 17 = wdExportFormatPDF
        if not Path(dst).exists():
            raise ConvertError("Word produced no PDF.")
        return [str(dst)]
    except ConvertError:
        Path(dst).unlink(missing_ok=True)
        raise
    except Exception as e:
        Path(dst).unlink(missing_ok=True)
        raise ConvertError(f"Word export failed: {e}")
    finally:
        try:
            if doc is not None:
                doc.Close(False)
        except Exception:
            pass
        try:
            if word is not None:
                word.Quit()
        except Exception:
            pass
        pythoncom.CoUninitialize()


def _office_pdf_soffice(src):
    """Export through LibreOffice headless (high fidelity, no MS Office needed)."""
    soffice = _bundled("soffice.exe") or _find_tool("soffice", (
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
    ))
    if not soffice:
        raise ConvertError("LibreOffice not found.")
    dst = unique_path(Path(src).with_suffix(".pdf"))
    with tempfile.TemporaryDirectory() as outdir:
        # A private profile dir lets us convert even if LibreOffice is open.
        prof = (Path(outdir) / "profile").as_uri()
        _run([soffice, "-env:UserInstallation=" + prof, "--headless",
              "--norestore", "--convert-to", "pdf:writer_pdf_Export",
              "--outdir", outdir, str(src)], "LibreOffice")
        produced = Path(outdir) / (Path(src).stem + ".pdf")
        if not produced.exists():
            pdfs = list(Path(outdir).glob("*.pdf"))
            if not pdfs:
                raise ConvertError("LibreOffice did not produce a PDF.")
            produced = pdfs[0]
        shutil.move(str(produced), str(dst))
    return [str(dst)]


def _doc_to_pdf_pandoc(src):
    """Route: pandoc -> standalone HTML -> Edge headless print-to-pdf.

    Used for markdown/html/txt/tex/epub, which have no rich source layout to
    preserve.
    """
    pd = _pandoc()
    edge = _edge()
    dst = unique_path(Path(src).with_suffix(".pdf"))
    with tempfile.TemporaryDirectory() as tmp:
        html = os.path.join(tmp, "out.html")
        css = os.path.join(tmp, "print.css")
        with open(css, "w", encoding="utf-8") as f:
            f.write("body{font-family:Segoe UI,Arial,sans-serif;"
                    "max-width:48em;margin:2em auto;line-height:1.5;}"
                    "pre{background:#f4f4f4;padding:1em;overflow-x:hidden;}")
        _run([pd, str(src)] + _pandoc_args(src) +
             ["-s", "--embed-resources", "-c", css, "-o", html], "Pandoc")
        _run_producing(
            [edge, "--headless=new", "--disable-gpu", "--no-pdf-header-footer",
             f"--user-data-dir={os.path.join(tmp, 'edgeprofile')}",
             f"--print-to-pdf={dst}", Path(html).as_uri()], "Edge", dst)
        if not Path(dst).exists():
            raise ConvertError("Edge did not produce a PDF.")
    return [str(dst)]


# ---------------------------------------------------------------- pdf

def convert_pdf(src, dst_ext):
    if dst_ext == "docx":
        try:
            from pdf2docx import Converter
        except ImportError:
            raise ConvertError("pdf2docx not installed. "
                               "Run: pip install --user pdf2docx")
        dst = unique_path(Path(src).with_suffix(".docx"))
        cv = Converter(str(src))
        try:
            cv.convert(str(dst))
        except Exception:
            Path(dst).unlink(missing_ok=True)
            raise
        finally:
            cv.close()
        return [str(dst)]

    try:
        import pymupdf as fitz
    except ImportError:
        try:
            import fitz  # older PyMuPDF name
        except ImportError:
            raise ConvertError("PyMuPDF not installed. "
                               "Run: pip install --user pymupdf")

    doc = fitz.open(str(src))
    try:
        if dst_ext == "txt":
            dst = unique_path(Path(src).with_suffix(".txt"))
            text = "\n".join(page.get_text() for page in doc)
            dst.write_text(text, encoding="utf-8")
            return [str(dst)]
        # png / jpg: render each page at 150 dpi
        outputs = []
        stem = Path(src).with_suffix("")
        for i, page in enumerate(doc, start=1):
            pix = page.get_pixmap(dpi=150)
            out = unique_path(Path(f"{stem}_p{i}.{dst_ext}"))
            pix.save(str(out))
            outputs.append(str(out))
        return outputs
    finally:
        doc.close()


# ---------------------------------------------------------------- archives

def _unique_arcname(name, used):
    """A top-level archive entry name not already present in `used`."""
    if name not in used:
        used.add(name)
        return name
    stem, ext = os.path.splitext(name)
    i = 1
    while f"{stem} ({i}){ext}" in used:
        i += 1
    new = f"{stem} ({i}){ext}"
    used.add(new)
    return new


def _add_path_to_zip(zf, path, top):
    p = Path(path)
    if p.is_dir():
        for child in sorted(p.rglob("*")):
            if child.is_file():
                rel = str(child.relative_to(p)).replace("\\", "/")
                zf.write(child, f"{top}/{rel}")
    else:
        zf.write(p, top)


def make_archive(paths, fmt, dst_dir=None):
    paths = [os.path.abspath(p) for p in paths]
    first = Path(paths[0])
    base_dir = Path(dst_dir) if dst_dir else first.parent
    if len(paths) == 1:
        # A folder has no extension to strip; a file does.
        stem = first.name if first.is_dir() else _split_name(first)[0]
    else:
        stem = first.parent.name or "archive"
    dst = unique_path(base_dir / f"{stem}.{fmt}")

    # Give each top-level item a distinct name so same-named files from
    # different folders don't collide inside the archive.
    used = set()
    entries = [(p, _unique_arcname(Path(p).name, used)) for p in paths]

    if fmt == "zip":
        with zipfile.ZipFile(dst, "w", zipfile.ZIP_DEFLATED) as zf:
            for p, top in entries:
                _add_path_to_zip(zf, p, top)
    elif fmt == "tar.gz":
        with tarfile.open(dst, "w:gz") as tf:
            for p, top in entries:
                tf.add(p, arcname=top)
    elif fmt == "7z":
        py7zr = _import_py7zr()
        with py7zr.SevenZipFile(dst, "w") as zf:
            for p, top in entries:
                zf.writeall(p, arcname=top)
    else:
        raise ConvertError(f"Unknown archive format: {fmt}")
    return [str(dst)]


def _extract_core(src, dest, ext):
    """Unpack `src` into the existing directory `dest`."""
    if ext == "zip":
        with zipfile.ZipFile(src) as zf:
            zf.extractall(dest)
    elif ext == "7z":
        py7zr = _import_py7zr()
        with py7zr.SevenZipFile(src) as zf:
            zf.extractall(dest)
    elif ext in ("tar", "tgz", "tar.gz", "tar.bz2", "tar.xz"):
        with tarfile.open(src) as tf:
            tf.extractall(dest, filter="data")
    elif ext in SINGLE_FILE_COMPRESSION:
        import bz2 as bz2mod
        import lzma
        opener = {"gz": gzip.open, "bz2": bz2mod.open, "xz": lzma.open}[ext]
        out = unique_path(Path(dest) / _split_name(src)[0])
        with opener(src, "rb") as f_in, open(out, "wb") as f_out:
            shutil.copyfileobj(f_in, f_out)
    elif ext == "rar":
        raise ConvertError("RAR extraction is not supported (proprietary "
                           "format). Use WinRAR or 7-Zip for .rar files.")
    else:
        raise ConvertError(f"Don't know how to extract .{ext}")


def extract_archive(src):
    src = Path(os.path.abspath(src))
    ext = ext_of(src)
    dest = unique_path(src.parent / _split_name(src)[0])
    made = False
    try:
        dest.mkdir(parents=True, exist_ok=True)
        made = True
        _extract_core(src, dest, ext)
    except Exception:
        # Don't leave an empty folder behind if the archive was bad.
        if made:
            try:
                if dest.is_dir() and not any(dest.iterdir()):
                    dest.rmdir()
            except OSError:
                pass
        raise
    return [str(dest)]


def repack_archive(src, fmt):
    """Extract an archive and repack its contents into a different format.

    Works for single-file .gz/.bz2/.xz too: the payload is decompressed first,
    then packed into the requested container.
    """
    src = os.path.abspath(src)
    ext = ext_of(src)
    with tempfile.TemporaryDirectory(dir=str(Path(src).parent)) as tmp:
        extracted = Path(tmp) / "extracted"
        extracted.mkdir()
        _extract_core(src, extracted, ext)
        items = [str(p) for p in extracted.iterdir()]
        if not items:
            raise ConvertError("Archive is empty.")
        dst = unique_path(Path(src).parent / f"{_split_name(src)[0]}.{fmt}")
        result = make_archive(items, fmt, dst_dir=tmp)
        shutil.move(result[0], dst)
    return [str(dst)]


# ---------------------------------------------------------------- dispatch

def convert_one(path, target_ext, quality=85):
    """Convert a single file; returns list of output paths."""
    path = os.path.abspath(path)  # tools like Edge resolve against their own CWD
    fam = family_of(path)
    if fam == "image":
        return convert_image(path, target_ext, quality)
    if fam in ("video", "audio"):
        return convert_av(path, target_ext)
    if fam == "document":
        return convert_doc(path, target_ext)
    if fam == "pdf":
        return convert_pdf(path, target_ext)
    if fam == "archive":
        return repack_archive(path, target_ext)
    raise ConvertError(f"No converter for {Path(path).name}")


def run_selftest():
    """Exercise every conversion family on scratch files.

    Returns a list of (name, status, detail). Used to verify a build (notably
    the standalone bundle, where a missing dependency would only show here).
    """
    import tempfile
    results = []
    tmp = tempfile.mkdtemp(prefix="qc-selftest-")

    def trial(name, fn):
        try:
            outs = fn()
            ok = bool(outs) and all(os.path.exists(o) for o in outs)
            results.append((name, "OK" if ok else "NO-OUTPUT", ""))
            return outs
        except Exception as e:
            results.append((name, "FAIL", f"{e.__class__.__name__}: {e}"))
            return []

    from PIL import Image
    png = os.path.join(tmp, "t.png")
    Image.new("RGB", (48, 48), (200, 40, 40)).save(png)

    trial("image: png->jpg (Pillow)", lambda: convert_image(png, "jpg"))
    trial("image: png->webp", lambda: convert_image(png, "webp"))
    trial("image: png->heic (pillow-heif)", lambda: convert_image(png, "heic"))
    pdfs = trial("image: png->pdf", lambda: convert_image(png, "pdf"))
    if pdfs:
        pdf = pdfs[0]
        trial("pdf: ->png (PyMuPDF)", lambda: convert_pdf(pdf, "png"))
        trial("pdf: ->txt (PyMuPDF)", lambda: convert_pdf(pdf, "txt"))
        trial("pdf: ->docx (pdf2docx/cv2)", lambda: convert_pdf(pdf, "docx"))
    try:
        from docx import Document
        dpath = os.path.join(tmp, "t.docx")
        _d = Document()
        _d.add_heading("Quick Convert self-test", 0)
        _d.add_paragraph("Formatting fidelity check.")
        _d.save(dpath)
        trial("docx: ->pdf (Word/LibreOffice)", lambda: convert_doc(dpath, "pdf"))
    except Exception as e:
        results.append(("docx: ->pdf", "FAIL", f"{e.__class__.__name__}: {e}"))
    trial("archive: ->zip", lambda: make_archive([png], "zip"))
    trial("archive: ->7z (py7zr)", lambda: make_archive([png], "7z"))
    trial("archive: ->tar.gz", lambda: make_archive([png], "tar.gz"))

    for tool, fn in (("ffmpeg (video/audio)", _ffmpeg),
                     ("pandoc (documents)", _pandoc),
                     ("edge (doc->pdf)", _edge)):
        try:
            results.append((f"tool: {tool}", "PRESENT", fn()))
        except Exception as e:
            results.append((f"tool: {tool}", "MISSING",
                            str(e).splitlines()[0]))

    try:
        import win32com.client  # noqa: F401
        import pythoncom  # noqa: F401
        results.append(("lib: pywin32 (Word docx->pdf)", "PRESENT", ""))
    except Exception as e:
        results.append(("lib: pywin32 (Word docx->pdf)", "MISSING",
                        str(e).splitlines()[0] if str(e) else ""))

    shutil.rmtree(tmp, ignore_errors=True)
    return results
